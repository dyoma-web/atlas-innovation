# -*- coding: utf-8 -*-
"""Genera 'Respuesta Comentarios Atlas.docx' — respuesta punto a punto a los
comentarios del cliente (Comentarios Atlas.docx, 2026-08-15)."""
import io, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
OUT = r'c:\wamp_3\www\atlas\Respuesta Comentarios Atlas.docx'

AZUL = RGBColor(0x0E, 0x44, 0x78)
VERDE = RGBColor(0x1F, 0x7A, 0x33)
NARANJA = RGBColor(0xA8, 0x62, 0x0C)
GRIS = RGBColor(0x55, 0x70, 0x7F)

ITEMS = [
    ("1. \u201cSpotlighted\u201d en lugar de \u201cSpotlight cases\u201d", "HECHO",
     ["Se cambió la etiqueta en toda la interfaz (título de sección, encabezados de grupo en \u201cAll cases\u201d, ficha de caso y filtro de visibilidad), en inglés y en español."],
     []),
    ("2. Nuevo texto de introducción de los casos Spotlighted", "HECHO",
     ["Se reemplazó el texto por el redactado entregado, palabra por palabra en inglés.",
      "Para la versión en español del sitio se generó la traducción correspondiente: \u201cVeinte iniciativas fueron cuidadosamente revisadas y seleccionadas como casos spotlighted por su enfoque único de la comunicación pública\u2026\u201d (sujeta a validación)."],
     []),
    ("3. Acortar títulos: VALUEbot y Voix des Survivants", "HECHO",
     ["\u201cVALUEbot : Una plataforma generadora de chatbots\u2026\u201d \u2192 \u201cVALUEbot\u201d.",
      "\u201cVoix des Survivants : une communication publique\u2026\u201d \u2192 \u201cVoix des Survivants\u201d.",
      "Se mantuvo la grafía oficial \u201cVALUEbot\u201d (el comentario decía \u201cValue Bot\u201d, pero el nombre registrado de la plataforma es VALUEbot)."],
     []),
    ("4. Quitar mayúsculas a \u201cECOS DE CIENAGA ORO\u201d", "HECHO (+ sugerencia aplicada)",
     ["Ahora aparece como \u201cEcos de Ciénaga de Oro\u201d (con la tilde y el \u201cde\u201d que usa su propia descripción).",
      "Sugerencia aplicada por coherencia: se normalizaron los otros 7 títulos que estaban todo en mayúsculas: Extra Favelas, Digital Justice and Human Rights in Angola, Social Media Gov, Decodificando Direitos e Deveres\u2026, Judiciário em Evidência, FCB Indoor y Herramientas para modernizar la escritura en la UNC\u2026"],
     ["Si prefieren conservar alguno en mayúsculas (p. ej. FCB INDOOR como marca), se revierte en un minuto."]),
    ("5. Imágenes faltantes: continuar con fondo azul", "CONFIRMADO — sin acción necesaria",
     ["Los tres casos sin material (Portal TCU, Les ondes de la parole publique, Ancla de Voces de Chinú) ya muestran la pieza de marca sobre fondo azul con la textura topográfica y el nombre del país."],
     ["Si en el futuro llega material de estos casos, se integra con el mismo pipeline en minutos."]),
    ("6. Brasil como caso aparte en el mapa", "HECHO",
     ["Brasil se pinta en el azul más oscuro y lleva un señalador con flecha: \u201cBrazil · 63 cases\u201d (clicable, filtra el repositorio).",
      "La escala de color del mapa ahora se calcula SOLO con los demás países (1 / 2 / 3\u20135 / 6+), de modo que las diferencias entre ellos vuelven a ser visibles y la comparación con Brasil deja de aplastar al resto.",
      "En el panel \u201cMost represented\u201d, Brasil sale de las barras y se muestra como tarjeta destacada con su cifra y una nota: \u201c63 cases \u2014 56% of the repository. Shown separately so differences among the other countries stay visible.\u201d Las barras comparan a los 8 países siguientes en su propia escala."],
     []),
    ("7. Categoría visible en \u201cAll cases\u201d (spotlighted y noteworthy)", "HECHO",
     ["Las mini-tarjetas Spotlighted de \u201cAll cases\u201d ahora muestran la píldora de categoría junto al país.",
      "Las filas Noteworthy incorporan una columna con la píldora de color de su categoría (en pantallas pequeñas se oculta por espacio; la categoría siempre está en la ficha)."],
     []),
    ("8. Explicación de los niveles Spotlighted / High visibility / Noteworthy", "HECHO — texto por validar",
     ["Se añadió una fila de definiciones bajo \u201cHow the cases are organised\u201d, con el punto de color de cada nivel:",
      "\u2022 Spotlighted \u2014 Selected for extended treatment: full narrative, signs of impact and supporting visual material.",
      "\u2022 High visibility \u2014 Full entry: short description, long-form case text and links.",
      "\u2022 Noteworthy \u2014 Brief entry: name, organisation, country and short description.",
      "La redacción se construyó a partir de la lógica de la Designers Guide (el texto exacto no existía escrito); versiones en español incluidas."],
     ["Validar la redacción de las tres definiciones (EN y ES). Nota: el nivel se llama \u201cHigh visibility\u201d, no \u201chighlight\u201d."]),
    ("9. Enlace web en los casos Noteworthy", "EN ESPERA — depende de insumo",
     ["Sin acción todavía: quedamos atentos a la tabla actualizada con los enlaces (hoy solo 49 de los 112 casos tienen enlace en el Excel)."],
     ["Recibir el Excel actualizado con la columna de enlaces.",
      "Decidir el comportamiento: recomendamos mantener el clic abriendo la ficha (conserva el aviso de idioma) y añadir un botón \u2197 directo al sitio en cada fila; la alternativa (clic directo al sitio externo) se pierde la ficha. Esperamos su preferencia."]),
    ("10. Logotipo de Catálise junto al del PNUD", "HECHO",
     ["Se integró el lockup blanco entregado (UNDP + catálise) en la franja azul superior de estadísticas, esquina superior derecha, replicando la composición de la imagen de referencia.",
      "Al estar siempre sobre el azul Deep Ocean (#0e4478), el contraste del blanco queda garantizado en escritorio y móvil, y también en la versión standalone (logo embebido)."],
     ["Cuando se haga el embebido final en Drupal, confirmar si el sitio anfitrión también mostrará logos propios, para no duplicarlos."]),
    ("11. Espacio para el PDF descargable al final", "EN ESPERA — depende de insumo",
     ["Sin acción todavía. El diseño del bloque de descarga (card al cierre del repositorio) está definido y se implementa en cuanto exista el archivo."],
     ["Recibir el PDF final publicable (el \u201cFinal Report EN\u201d que tenemos fue entregado como referencia de diseño; entendemos que la versión pública puede diferir)."]),
    ("12. Video de Catálise al inicio", "EN ESPERA — depende de insumo",
     ["Sin acción todavía. Está previsto un bloque de video al inicio del repositorio."],
     ["Recibir el archivo (está en SharePoint y en borrador; David está gestionando el acceso).",
      "Definir el hosting: recomendamos servirlo como archivo hospedado (Drupal/YouTube/Vimeo) y no embebido, para no inflar la versión standalone."]),
]


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    h = doc.add_heading('', level=0)
    r = h.add_run('Atlas of Innovation — Respuesta a los comentarios')
    r.font.color.rgb = AZUL
    p = doc.add_paragraph()
    r = p.add_run('Referencia: "Comentarios Atlas.docx" · Respuesta del 15 de agosto de 2026\n'
                  'Versión publicada con los cambios: https://dyoma-web.github.io/atlas-innovation/  ·  '
                  'Se adjunta también el archivo standalone actualizado (funciona con doble clic, sin internet).')
    r.font.color.rgb = GRIS
    r.font.size = Pt(9.5)

    resumen = doc.add_paragraph()
    r = resumen.add_run('Resumen: 8 solicitudes implementadas y publicadas · 1 confirmación sin acción · '
                        '3 en espera de insumos (enlaces Noteworthy, PDF final, video).')
    r.bold = True
    r.font.color.rgb = AZUL

    for titulo, estado, acciones, pendientes in ITEMS:
        h = doc.add_heading('', level=2)
        r = h.add_run(titulo)
        r.font.color.rgb = AZUL
        r.font.size = Pt(12)

        p = doc.add_paragraph()
        r = p.add_run('Estado: ')
        r.bold = True
        r2 = p.add_run(estado)
        r2.bold = True
        r2.font.color.rgb = VERDE if estado.startswith(('HECHO', 'CONFIRMADO')) else NARANJA

        if acciones:
            p = doc.add_paragraph()
            p.add_run('Acciones emprendidas:').bold = True
            for a in acciones:
                doc.add_paragraph(a, style='List Bullet')
        if pendientes:
            p = doc.add_paragraph()
            p.add_run('Por resolver / en espera:').bold = True
            for x in pendientes:
                doc.add_paragraph(x, style='List Bullet')

    doc.add_heading('Nota final', level=2).runs[0].font.color.rgb = AZUL
    doc.add_paragraph(
        'Todos los cambios están desplegados en la versión en línea y en el standalone. '
        'Los tres puntos en espera se implementan tan pronto lleguen los insumos: '
        'la estructura del desarrollo ya los contempla, por lo que su integración es cuestión de horas, no de días.')

    doc.save(OUT)
    print('Escrito:', OUT)


if __name__ == '__main__':
    main()
